"""OneDrive adapter — ingests documents from a local OneDrive directory.

Source: local OneDrive root directory (F:\\OneDrive\\ post-2026-05-13 reorg).
Walks Outputs/, Reference/, Records/ top-level pillars. Skips binary
extensions and files >20MB.

Reference/ has a body-extract allowlist: active-pursuit subdirs get full
body extraction; everything else gets metadata-only rows (subject +
file_path, is_bulk=1).

Each file becomes a schema_type='DigitalDocument' row in the documents
typed table.
"""

from __future__ import annotations

import hashlib
import io
import json
from collections.abc import Iterator
from datetime import UTC, datetime
from pathlib import Path

from phdb.adapters.base import Adapter, AdapterRow, DedupStrategy
from phdb.log import get_logger

log = get_logger("phdb.adapters.onedrive")

MAX_BODY_LEN = 200_000
MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024

INCLUDE_TOP_DIRS = {"Outputs", "Reference", "Records"}

REFERENCE_BODY_ALLOWLIST = {
    "Mind Tools", "Weight Training Programs", "Weight Training Theory",
    "Nutrition", "Cooking", "Personal Development", "Spirituality",
    "Philosophy", "Mental Health", "Entrepreneurship", "ADHD",
    "Fat Loss", "Marek Health Physical", "Medicaid", "MMA", "Rugby",
    "Career Choice", "Sexuality", "Kink", "Supplements",
    "People",
}

TEXT_EXTENSIONS = {
    ".docx", ".pdf", ".txt", ".md", ".html", ".htm",
    ".ipynb", ".xlsx", ".xls", ".csv", ".json", ".rtf",
}

SKIP_EXTENSIONS = {
    ".jpg", ".jpeg", ".png", ".heic", ".gif", ".bmp", ".webp", ".svg", ".ico",
    ".mp3", ".wav", ".mp4", ".mov", ".m4a", ".aac", ".flac",
    ".epub", ".mobi", ".azw", ".azw3", ".opf",
    ".py", ".java", ".lua", ".php", ".js", ".ts",
    ".exe", ".dll", ".apk", ".dmg", ".zip", ".gz", ".tar",
    ".dds", ".tga", ".blp",
    ".db", ".sqlite", ".sqlite3",
}


# ---- Text extractors (duplicated from google_drive for adapter independence) --

def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]

        d = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in d.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in d.tables:
            for row in table.rows:
                row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_txt:
                    parts.append(row_txt)
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[docx-extract-unavailable: python-docx not installed]"
    except Exception as e:
        return f"[docx-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]

        r = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in r.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[pdf-extract-unavailable: pypdf not installed]"
    except Exception as e:
        return f"[pdf-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore[import-untyped]

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                row_txt = "\t".join(str(v) for v in row if v is not None)
                if row_txt.strip():
                    parts.append(row_txt)
        wb.close()
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[xlsx-extract-unavailable: openpyxl not installed]"
    except Exception as e:
        return f"[xlsx-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def _extract_ipynb(data: bytes) -> str:
    try:
        nb = json.loads(data)
        parts: list[str] = []
        for cell in nb.get("cells", []):
            src = "".join(cell.get("source", []))
            if src.strip():
                parts.append(src)
        return "\n\n".join(parts)[:MAX_BODY_LEN]
    except Exception as e:
        return f"[ipynb-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def _extract_text(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(enc)[:MAX_BODY_LEN]
        except (UnicodeDecodeError, ValueError):
            pass
    return data.decode("latin-1", errors="replace")[:MAX_BODY_LEN]


_EXTRACTORS: dict[str, tuple[str, object]] = {
    ".docx": ("docx", _extract_docx),
    ".pdf": ("pdf", _extract_pdf),
    ".xlsx": ("xlsx", _extract_xlsx),
    ".xls": ("xlsx", _extract_xlsx),
    ".ipynb": ("ipynb", _extract_ipynb),
    ".txt": ("plaintext", _extract_text),
    ".md": ("plaintext", _extract_text),
    ".csv": ("plaintext", _extract_text),
    ".json": ("plaintext", _extract_text),
    ".html": ("plaintext", _extract_text),
    ".htm": ("plaintext", _extract_text),
    ".rtf": ("plaintext", _extract_text),
}


def _derive_bucket(rel_parts: tuple[str, ...]) -> str:
    if not rel_parts:
        return "(root)"
    if len(rel_parts) == 1:
        return rel_parts[0]
    return f"{rel_parts[0]}/{rel_parts[1]}"


def _is_reference_body_allowed(rel_parts: tuple[str, ...]) -> bool:
    """Check if a file under Reference/ is in the body-extract allowlist."""
    if len(rel_parts) < 2 or rel_parts[0] != "Reference":
        return True
    return rel_parts[1] in REFERENCE_BODY_ALLOWLIST


class OneDriveAdapter(Adapter):
    """Ingest OneDrive documents from a local directory."""

    name = "onedrive"
    source_kind = "onedrive"
    file_kind = "local-files"
    schema_type = "DigitalDocument"
    target_table = "documents"
    dedup_strategy = DedupStrategy.RFC822_MESSAGE_ID
    batch_size = 500

    def compute_raw_hash(self, row: AdapterRow) -> str:
        return hashlib.sha256(
            f"onedrive|{row.rfc822_message_id or ''}".encode()
        ).hexdigest()

    def iter_rows(self, source_path: Path, **kwargs: object) -> Iterator[AdapterRow]:
        candidate_dirs = [
            source_path / d
            for d in sorted(INCLUDE_TOP_DIRS)
            if (source_path / d).is_dir()
        ]

        for top_dir in candidate_dirs:
            for fpath in sorted(top_dir.rglob("*")):
                if not fpath.is_file():
                    continue

                suffix = fpath.suffix.lower()
                if suffix in SKIP_EXTENSIONS:
                    continue
                if suffix not in TEXT_EXTENSIONS:
                    continue

                try:
                    stat = fpath.stat()
                    if stat.st_size > MAX_FILE_SIZE_BYTES:
                        continue
                except OSError:
                    continue

                rel_parts = fpath.relative_to(source_path).parts
                bucket = _derive_bucket(rel_parts)
                relpath = str(fpath.relative_to(source_path))
                path_hash = hashlib.sha1(str(fpath).encode()).hexdigest()[:16]
                msg_id = f"onedrive:{path_hash}"

                try:
                    mtime = datetime.fromtimestamp(
                        stat.st_mtime, tz=UTC
                    ).isoformat()
                except OSError:
                    mtime = None

                try:
                    ctime = datetime.fromtimestamp(
                        stat.st_ctime, tz=UTC
                    ).isoformat()
                except OSError:
                    ctime = None

                raw_hash = hashlib.sha256(
                    f"onedrive|{msg_id}".encode()
                ).hexdigest()

                is_bulk = 1 if rel_parts[0] == "Reference" else 0

                if _is_reference_body_allowed(rel_parts):
                    entry = _EXTRACTORS.get(suffix)
                    if not entry:
                        continue
                    body_source, extractor = entry
                    try:
                        data = fpath.read_bytes()
                    except Exception:
                        continue
                    body = extractor(data)  # type: ignore[operator]
                    if not body or not body.strip():
                        continue
                    body = body[:MAX_BODY_LEN]
                else:
                    body = None
                    body_source = None

                yield AdapterRow(
                    schema_type="DigitalDocument",
                    rfc822_message_id=msg_id,
                    subject=fpath.name,
                    date_sent=mtime,
                    body_text=body,
                    body_text_source=str(body_source) if body_source else None,
                    is_bulk=is_bulk,
                    raw_hash=raw_hash,
                    body_text_hash=hashlib.sha256(body.encode()).hexdigest() if body else None,
                    file_path=relpath,
                    file_size=stat.st_size,
                    ctime=ctime,
                    bucket=bucket,
                )
