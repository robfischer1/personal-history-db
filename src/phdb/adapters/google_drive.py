"""Google Drive adapter — ingests documents from Google Takeout zips.

Source: a Google Takeout zip (or extracted directory) containing ``Drive/``
paths with text-bearing documents.

Each extracted file becomes a schema_type='DigitalDocument' row.
Text extraction covers docx, pdf, xlsx, ipynb, html, csv, json, txt, md, rtf.
External-library formats (docx, pdf, xlsx, html) gracefully degrade when
their dependencies are not installed.
"""

from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import zipfile
from collections.abc import Iterator
from datetime import datetime
from pathlib import Path

from phdb.adapters.base import Adapter, AdapterRow, DedupStrategy
from phdb.log import get_logger

log = get_logger("phdb.adapters.google_drive")

MAX_BODY_LEN = 200_000

TEXT_EXTENSIONS = {
    ".docx", ".pdf", ".txt", ".md", ".html", ".htm",
    ".ipynb", ".xlsx", ".xls", ".csv", ".json", ".rtf",
}

SKIP_EXTENSIONS = {
    ".feather", ".gz", ".tar", ".zip", ".bz2", ".xz",
    ".jpg", ".jpeg", ".png", ".heic", ".gif", ".ico", ".svg", ".bmp", ".webp",
    ".mp3", ".wav", ".mov", ".mp4", ".m4a", ".aac",
    ".vcf", ".kmz", ".kml", ".properties", ".apk", ".exe", ".dll", ".dmg",
    ".parquet",
}

SKIP_PATH_PATTERNS = [
    re.compile(r"TitaniumBackup"),
    re.compile(r"\.gdrive$"),
    re.compile(r"/Trash/"),
    re.compile(r"/Copy of [^/]+$"),
]


# ---- Text extractors -------------------------------------------------------

def extract_docx(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]

        d = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in d.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in d.tables:
            for row in table.rows:
                row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_txt:
                    parts.append(row_txt)
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[docx-extract-unavailable: python-docx not installed]"
    except Exception as e:
        return f"[docx-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]

        r = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in r.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
            if sum(len(p) for p in parts) > MAX_BODY_LEN:
                break
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[pdf-extract-unavailable: pypdf not installed]"
    except Exception as e:
        return f"[pdf-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore[import-untyped]

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    parts.append(" | ".join(cells))
                if sum(len(p) for p in parts) > MAX_BODY_LEN:
                    wb.close()
                    return "\n".join(parts)[:MAX_BODY_LEN]
        wb.close()
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[xlsx-extract-unavailable: openpyxl not installed]"
    except Exception as e:
        return f"[xlsx-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_html(data: bytes) -> str:
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(data, "lxml")
        return soup.get_text("\n", strip=True)[:MAX_BODY_LEN]
    except ImportError:
        return "[html-extract-unavailable: bs4 not installed]"
    except Exception as e:
        return f"[html-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_ipynb(data: bytes) -> str:
    try:
        nb = json.loads(data)
        parts: list[str] = []
        for cell in nb.get("cells", []):
            ctype = cell.get("cell_type", "")
            src = cell.get("source", "")
            if isinstance(src, list):
                src = "".join(src)
            if not src.strip():
                continue
            parts.append(f"[{ctype}]\n{src}")
            if sum(len(p) for p in parts) > MAX_BODY_LEN:
                break
        return "\n\n".join(parts)[:MAX_BODY_LEN]
    except Exception as e:
        return f"[ipynb-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_csv(data: bytes) -> str:
    try:
        text = data.decode("utf-8-sig", errors="replace")
        rdr = csv.reader(io.StringIO(text))
        parts: list[str] = []
        for row in rdr:
            if any(c.strip() for c in row):
                parts.append(" | ".join(row))
            if sum(len(p) for p in parts) > MAX_BODY_LEN:
                break
        return "\n".join(parts)[:MAX_BODY_LEN]
    except Exception as e:
        return f"[csv-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_json(data: bytes) -> str:
    try:
        obj = json.loads(data)
        return json.dumps(obj, indent=2)[:MAX_BODY_LEN]
    except Exception:
        return data.decode("utf-8", errors="replace")[:MAX_BODY_LEN]


def extract_txt(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")[:MAX_BODY_LEN]


def extract_rtf(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    text = re.sub(r"\\[a-zA-Z]+\d*\s?", "", text)
    text = re.sub(r"[{}\\]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:MAX_BODY_LEN]


EXTRACTORS: dict[str, tuple[str, object]] = {
    ".docx": ("docx", extract_docx),
    ".pdf": ("pdf", extract_pdf),
    ".xlsx": ("xlsx", extract_xlsx),
    ".xls": ("xlsx", extract_xlsx),
    ".ipynb": ("ipynb", extract_ipynb),
    ".html": ("html", extract_html),
    ".htm": ("html", extract_html),
    ".csv": ("csv", extract_csv),
    ".json": ("json", extract_json),
    ".txt": ("txt", extract_txt),
    ".md": ("txt", extract_txt),
    ".rtf": ("rtf", extract_rtf),
}


def derive_bucket(relpath: str) -> str:
    """First 2 path components after Drive/."""
    parts = relpath.split("/")
    while parts and parts[0] in ("Takeout", "Drive"):
        parts = parts[1:]
    if len(parts) <= 1:
        return "(root)"
    if parts[0].startswith(("00 ", "01 ", "02 ", "03 ", "04 ")) and len(parts) >= 2:
        return "/".join(parts[:2])
    return parts[0]


def _collect_docx_stems(names: list[str]) -> set[tuple[str, str]]:
    stems: set[tuple[str, str]] = set()
    for n in names:
        p = Path(n)
        if p.suffix.lower() == ".docx":
            stems.add((str(p.parent), p.stem))
    return stems


def _is_pdf_shadowed_by_docx(name: str, docx_stems: set[tuple[str, str]]) -> bool:
    p = Path(name)
    if p.suffix.lower() != ".pdf":
        return False
    return (str(p.parent), p.stem) in docx_stems


def _passes_skip_filters(name: str) -> bool:
    if any(p.search(name) for p in SKIP_PATH_PATTERNS):
        return False
    ext = Path(name).suffix.lower()
    if ext in SKIP_EXTENSIONS:
        return False
    return ext in TEXT_EXTENSIONS


class GoogleDriveAdapter(Adapter):
    """Ingest Google Drive documents from Takeout zips or directories."""

    name = "google_drive"
    source_kind = "google-drive"
    file_kind = "zip"
    schema_type = "DigitalDocument"
    dedup_strategy = DedupStrategy.CONTENT_HASH
    batch_size = 500

    def compute_raw_hash(self, row: AdapterRow) -> str:
        seed = f"google-drive|{row.extra.get('relpath', '')}|{len(row.body_text or '')}"
        return hashlib.sha256(seed.encode("utf-8")).hexdigest()

    def iter_rows(self, source_path: Path, **kwargs: object) -> Iterator[AdapterRow]:
        if source_path.is_file() and source_path.suffix == ".zip":
            yield from self._iter_zip(source_path)
        elif source_path.is_dir():
            yield from self._iter_dir(source_path)

    def _iter_zip(self, zip_path: Path) -> Iterator[AdapterRow]:
        with zipfile.ZipFile(zip_path) as zf:
            candidates: list[str] = []
            for info in zf.infolist():
                if info.is_dir():
                    continue
                name = info.filename
                if "Drive/" not in name:
                    continue
                if _passes_skip_filters(name):
                    candidates.append(name)

            docx_stems = _collect_docx_stems(candidates)

            for fi, name in enumerate(sorted(candidates)):
                if _is_pdf_shadowed_by_docx(name, docx_stems):
                    continue
                ext = Path(name).suffix.lower()
                entry = EXTRACTORS.get(ext)
                if not entry:
                    continue
                body_source, extractor = entry
                try:
                    data = zf.read(name)
                except Exception:
                    continue
                info = zf.getinfo(name)
                mtime = datetime(*info.date_time).isoformat() if info.date_time else None
                body = extractor(data)  # type: ignore[operator]
                if not body or not body.strip():
                    continue
                body = body[:MAX_BODY_LEN]
                bucket = derive_bucket(name)
                filename = Path(name).name
                raw_hash = hashlib.sha256(
                    f"google-drive|{name}|{len(body)}".encode()
                ).hexdigest()

                yield AdapterRow(
                    schema_type="DigitalDocument",
                    rfc822_message_id=f"google-drive:{raw_hash}",
                    subject=filename[:200],
                    sender_address=self.owner_sender("google-drive")[0],
                    sender_name=bucket,
                    direction="self",
                    date_sent=mtime,
                    body_text=body,
                    body_text_source=str(body_source),
                    is_bulk=0,
                    source_byte_offset=fi,
                    source_byte_length=len(body),
                    raw_hash=raw_hash,
                    body_text_hash=hashlib.sha256(body.encode()).hexdigest(),
                    thread_key=f"google-drive:{bucket}",
                    extra={"relpath": name},
                )

    def _iter_dir(self, dir_path: Path) -> Iterator[AdapterRow]:
        candidates: list[str] = []
        for p in sorted(dir_path.rglob("*")):
            if not p.is_file():
                continue
            relstr = str(p.relative_to(dir_path)).replace("\\", "/")
            if _passes_skip_filters(relstr):
                candidates.append(relstr)

        docx_stems = _collect_docx_stems(candidates)

        for fi, relstr in enumerate(sorted(candidates)):
            if _is_pdf_shadowed_by_docx(relstr, docx_stems):
                continue
            ext = Path(relstr).suffix.lower()
            entry = EXTRACTORS.get(ext)
            if not entry:
                continue
            body_source, extractor = entry
            full_path = dir_path / relstr
            try:
                data = full_path.read_bytes()
            except Exception:
                continue
            mtime = datetime.fromtimestamp(full_path.stat().st_mtime).isoformat()
            body = extractor(data)  # type: ignore[operator]
            if not body or not body.strip():
                continue
            body = body[:MAX_BODY_LEN]
            bucket = derive_bucket(relstr)
            filename = Path(relstr).name
            raw_hash = hashlib.sha256(
                f"google-drive|{relstr}|{len(body)}".encode()
            ).hexdigest()

            yield AdapterRow(
                schema_type="DigitalDocument",
                rfc822_message_id=f"google-drive:{raw_hash}",
                subject=filename[:200],
                sender_address=self.owner_sender("google-drive")[0],
                sender_name=bucket,
                direction="self",
                date_sent=mtime,
                body_text=body,
                body_text_source=str(body_source),
                is_bulk=0,
                source_byte_offset=fi,
                source_byte_length=len(body),
                raw_hash=raw_hash,
                body_text_hash=hashlib.sha256(body.encode()).hexdigest(),
                thread_key=f"google-drive:{bucket}",
                extra={"relpath": relstr},
            )
