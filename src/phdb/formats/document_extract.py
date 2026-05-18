"""Shared text extraction functions for document-bearing adapters.

Extracts text from common file formats (docx, pdf, xlsx, ipynb, html, csv,
json, txt, rtf). External-library formats gracefully degrade when their
dependencies are not installed.

Used by: google_drive_zip format parser, onedrive adapter, and any future
adapter that needs to pull text from document files.
"""

from __future__ import annotations

import csv
import io
import json
import re
from collections.abc import Callable

MAX_BODY_LEN = 200_000


def extract_docx(data: bytes) -> str:
    """Extract text from a .docx file (paragraphs + tables)."""
    try:
        from docx import Document  # type: ignore[import-not-found]

        d = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in d.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in d.tables:
            for row in table.rows:
                row_txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_txt:
                    parts.append(row_txt)
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[docx-extract-unavailable: python-docx not installed]"
    except Exception as e:
        return f"[docx-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_pdf(data: bytes) -> str:
    """Extract text from a PDF file page by page."""
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]

        r = PdfReader(io.BytesIO(data))
        parts: list[str] = []
        for page in r.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
            if sum(len(p) for p in parts) > MAX_BODY_LEN:
                break
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[pdf-extract-unavailable: pypdf not installed]"
    except Exception as e:
        return f"[pdf-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_xlsx(data: bytes) -> str:
    """Extract text from an Excel workbook, iterating by sheet name."""
    try:
        from openpyxl import load_workbook  # type: ignore[import-untyped]

        wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    parts.append(" | ".join(cells))
                if sum(len(p) for p in parts) > MAX_BODY_LEN:
                    wb.close()
                    return "\n".join(parts)[:MAX_BODY_LEN]
        wb.close()
        return "\n".join(parts)[:MAX_BODY_LEN]
    except ImportError:
        return "[xlsx-extract-unavailable: openpyxl not installed]"
    except Exception as e:
        return f"[xlsx-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_html(data: bytes) -> str:
    """Extract text from HTML using BeautifulSoup."""
    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(data, "lxml")
        return soup.get_text("\n", strip=True)[:MAX_BODY_LEN]
    except ImportError:
        return "[html-extract-unavailable: bs4 not installed]"
    except Exception as e:
        return f"[html-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_ipynb(data: bytes) -> str:
    """Extract text from a Jupyter notebook (all cell sources)."""
    try:
        nb = json.loads(data)
        parts: list[str] = []
        for cell in nb.get("cells", []):
            ctype = cell.get("cell_type", "")
            src = cell.get("source", "")
            if isinstance(src, list):
                src = "".join(src)
            if not src.strip():
                continue
            parts.append(f"[{ctype}]\n{src}")
            if sum(len(p) for p in parts) > MAX_BODY_LEN:
                break
        return "\n\n".join(parts)[:MAX_BODY_LEN]
    except Exception as e:
        return f"[ipynb-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_csv(data: bytes) -> str:
    """Extract text from a CSV file as pipe-delimited rows."""
    try:
        text = data.decode("utf-8-sig", errors="replace")
        rdr = csv.reader(io.StringIO(text))
        parts: list[str] = []
        for row in rdr:
            if any(c.strip() for c in row):
                parts.append(" | ".join(row))
            if sum(len(p) for p in parts) > MAX_BODY_LEN:
                break
        return "\n".join(parts)[:MAX_BODY_LEN]
    except Exception as e:
        return f"[csv-extract-failed: {type(e).__name__}: {str(e)[:200]}]"


def extract_json(data: bytes) -> str:
    """Extract text from JSON (pretty-printed)."""
    try:
        obj = json.loads(data)
        return json.dumps(obj, indent=2)[:MAX_BODY_LEN]
    except Exception:
        return data.decode("utf-8", errors="replace")[:MAX_BODY_LEN]


def extract_txt(data: bytes) -> str:
    """Extract plain text (UTF-8 with fallback)."""
    return data.decode("utf-8", errors="replace")[:MAX_BODY_LEN]


def extract_rtf(data: bytes) -> str:
    """Extract text from RTF by stripping control words."""
    text = data.decode("utf-8", errors="replace")
    text = re.sub(r"\\[a-zA-Z]+\d*\s?", "", text)
    text = re.sub(r"[{}\\]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:MAX_BODY_LEN]


EXTRACTORS: dict[str, tuple[str, Callable[[bytes], str]]] = {
    ".docx": ("docx", extract_docx),
    ".pdf": ("pdf", extract_pdf),
    ".xlsx": ("xlsx", extract_xlsx),
    ".xls": ("xlsx", extract_xlsx),
    ".ipynb": ("ipynb", extract_ipynb),
    ".html": ("html", extract_html),
    ".htm": ("html", extract_html),
    ".csv": ("csv", extract_csv),
    ".json": ("json", extract_json),
    ".txt": ("txt", extract_txt),
    ".md": ("txt", extract_txt),
    ".rtf": ("rtf", extract_rtf),
}
